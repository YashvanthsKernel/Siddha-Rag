"""
=================================================================================
SIDDHA RAG PIPELINE - MAIN ORCHESTRATOR
=================================================================================

This is the main pipeline that processes documents through all stages:
1. Document Ingestion (PDF/DOCX reading)
2. Text Cleaning (removing noise, formatting)
3. Text Chunking (splitting into manageable pieces)
4. Embedding Generation (converting to vectors)
5. Vector Storage (storing in ChromaDB)

Each step is documented with detailed explanations.
=================================================================================
"""

import os
import sys
import logging
from pathlib import Path
from typing import List, Dict

# Add project root to Python path for imports
PROJECT_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(PROJECT_ROOT))

# Configure logging for better visibility
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Import pipeline components
try:
    from PyPDF2 import PdfReader
except ImportError:
    logging.warning("PyPDF2 not installed. PDF reading may fail. Install: pip install PyPDF2")

# OCR imports for scanned PDF fallback
# DISABLED: Set to False to skip OCR (already processed separately)
OCR_AVAILABLE = False
# try:
#     import pytesseract
#     from pdf2image import convert_from_path
#     OCR_AVAILABLE = True
#     logging.info("✅ OCR support available (pytesseract + pdf2image)")
# except ImportError:
#     logging.info("ℹ️  OCR not available. Scanned PDFs won't be processed. Install: pip install pytesseract pdf2image")
logging.info("ℹ️  OCR disabled (scanned PDFs already processed separately)")

from docx import Document
from src.pipeline.cleaner import TextCleaner
from src.pipeline.chunking import TextChunker
from src.rag.embeddings import EmbeddingManager


# =============================================================================
# STEP 1: DOCUMENT INGESTION
# =============================================================================

class DocumentIngester:
    """
    Handles reading documents from various formats.
    
    Purpose:
        - Read PDF files using PyPDF2 (no poppler required)
        - Fall back to OCR for scanned PDFs (requires Tesseract + Poppler)
        - Read DOCX files using python-docx
        - Extract plain text from documents
        - Handle errors gracefully
    
    Methods:
        - read_pdf(): Extract text from PDF (with OCR fallback)
        - read_docx(): Extract text from DOCX
        - process_all_documents(): Process entire directory
    """
    
    def __init__(self, raw_data_dir: str = "data/raw"):
        """
        Initialize the document ingester.
        
        Args:
            raw_data_dir: Directory containing raw PDF/DOCX files
        """
        self.raw_data_dir = Path(raw_data_dir)
        logging.info(f"📂 Initialized DocumentIngester")
        logging.info(f"   Source directory: {self.raw_data_dir}")
        if OCR_AVAILABLE:
            logging.info(f"   OCR fallback: ENABLED")
        else:
            logging.info(f"   OCR fallback: DISABLED (install pytesseract + pdf2image)")
    
    def read_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file using PyPDF2, with OCR fallback.
        
        Strategy:
            1. First try PyPDF2 (fast, works for text-based PDFs)
            2. If empty result, fall back to OCR (for scanned PDFs)
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        try:
            logging.info(f"   📄 Reading PDF: {file_path.name}")
            
            # Initialize PDF reader
            reader = PdfReader(str(file_path))
            
            # Handle encrypted PDFs
            if reader.is_encrypted:
                try:
                    if reader.decrypt("") == 0:
                        logging.warning(f"      ⚠️  PDF is password-protected: {file_path.name}")
                        return ""
                except Exception as decrypt_err:
                    logging.warning(f"      ⚠️  Cannot decrypt PDF: {file_path.name}")
                    return ""
            
            text = ""
            
            # Extract text from each page using PyPDF2
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as page_err:
                    continue
            
            pages = len(reader.pages)
            chars = len(text.strip())
            
            # If PyPDF2 returned empty but OCR is available, try OCR
            if chars == 0 and OCR_AVAILABLE and pages > 0:
                logging.info(f"      🔄 PyPDF2 returned empty, trying OCR...")
                text = self._ocr_pdf(file_path)
                if text:
                    logging.info(f"      ✅ OCR extracted {len(text):,} characters from {pages} pages")
                    return text.strip()
                else:
                    logging.warning(f"      ⚠️  OCR also failed for {file_path.name}")
                    return ""
            elif chars == 0:
                logging.warning(f"      ⚠️  Extracted 0 characters from {pages} pages (scanned PDF?)")
                if not OCR_AVAILABLE:
                    logging.warning(f"         💡 Enable OCR: pip install pytesseract pdf2image")
                return ""
            
            logging.info(f"      ✅ Extracted {chars:,} characters from {pages} pages")
            return text.strip()
            
        except Exception as e:
            error_msg = str(e)
            if "PyCryptodome" in error_msg or "AES" in error_msg:
                logging.warning(f"      ⚠️  This PDF is encrypted. Try: pip install pycryptodome")
            else:
                logging.error(f"      ❌ Error reading PDF: {e}")
            return ""
    
    def _ocr_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF using OCR (Tesseract + pdf2image).
        
        This is used as a fallback for scanned PDFs that PyPDF2 can't process.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        if not OCR_AVAILABLE:
            return ""
        
        try:
            # Convert PDF pages to images
            images = convert_from_path(str(file_path))
            text = ""
            
            for i, image in enumerate(images):
                # OCR each page (supports Tamil + English)
                page_text = pytesseract.image_to_string(image, lang='eng+tam')
                text += page_text + "\n\n"
            
            return text.strip()
            
        except Exception as e:
            logging.error(f"      ❌ OCR error: {e}")
            return ""
    
    def read_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Why python-docx?
            - Native support for DOCX format
            - Preserves paragraph structure
            - Reliable and well-maintained
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        
        Process:
            1. Open DOCX with Document()
            2. Extract all paragraphs
            3. Join with newlines
        """
        try:
            logging.info(f"   📝 Reading DOCX: {file_path.name}")
            
            # Load document
            doc = Document(str(file_path))
            
            # Extract all paragraph text
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)
            
            # Log success
            logging.info(f"      ✅ Extracted {len(text):,} characters from {len(paragraphs)} paragraphs")
            
            return text
            
        except Exception as e:
            logging.error(f"      ❌ Error reading DOCX: {e}")
            return ""
    
    def process_all_documents(self) -> List[Dict]:
        """
        Process all PDF and DOCX files in the raw data directory.
        
        Returns:
            List of dictionaries, each containing:
                - filename: Original filename
                - text: Extracted text content
                - source: Full file path
                - file_type: 'pdf' or 'docx'
                - char_count: Number of characters
        
        Process:
            1. Scan directory for PDF/DOCX files
            2. Process each file based on extension
            3. Collect results with metadata
            4. Return formatted list
        """
        documents = []
        
        # Find all PDF and DOCX files
        pdf_files = list(self.raw_data_dir.glob("*.pdf"))
        docx_files = list(self.raw_data_dir.glob("*.docx"))
        all_files = pdf_files + docx_files
        
        if not all_files:
            logging.warning(f"⚠️  No PDF or DOCX files found in {self.raw_data_dir}")
            return documents
        
        logging.info(f"📖 Found {len(all_files)} documents to process")
        logging.info(f"   PDFs: {len(pdf_files)}, DOCX: {len(docx_files)}")
        
        # Process each file
        for file_path in all_files:
            filename = file_path.name
            file_ext = file_path.suffix.lower()
            
            # Read based on file type
            if file_ext == '.pdf':
                text = self.read_pdf(file_path)
                file_type = 'pdf'
            elif file_ext == '.docx':
                text = self.read_docx(file_path)
                file_type = 'docx'
            else:
                continue
            
            # Skip if no text extracted
            if not text or len(text.strip()) == 0:
                logging.warning(f"      ⚠️  Empty content from {filename}")
                continue
            
            # Add to documents list with metadata
            documents.append({
                'filename': filename,
                'text': text,
                'source': str(file_path),
                'file_type': file_type,
                'char_count': len(text)
            })
        
        logging.info(f"✅ Successfully processed {len(documents)} documents")
        return documents


# =============================================================================
# MAIN PIPELINE ORCHESTRATION
# =============================================================================

def run_complete_pipeline(
    raw_data_dir: str = "data/raw",
    processed_dir: str = "data/processed",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    reset_vectordb: bool = False
) -> int:
    """
    Execute the complete RAG pipeline from raw documents to vector database.
    
    Pipeline Stages:
        1. DOCUMENT INGESTION - Read PDF/DOCX files
        2. TEXT CLEANING - Remove noise, normalize formatting
        3. TEXT CHUNKING - Split into manageable pieces
        4. EMBEDDING GENERATION - Convert to vector representations
        5. VECTOR STORAGE - Store in ChromaDB for similarity search
    
    Args:
        raw_data_dir: Directory with raw PDF/DOCX files
        processed_dir: Directory for intermediate outputs
        chunk_size: Maximum characters per chunk (default: 1000)
        chunk_overlap: Characters overlap between chunks (default: 200)
        reset_vectordb: If True, deletes existing vector DB
    
    Returns:
        int: Number of chunks successfully processed
    
    Why these defaults?
        - chunk_size=1000: Good balance between context and granularity
        - chunk_overlap=200: Ensures context continuity across chunks
"""
    
    # Print pipeline header
    print("="*80)
    print("🚀 SIDDHA RAG PIPELINE - STARTING")
    print("="*80)
    print("\n📋 Configuration:")
    print(f"   Raw data source: {raw_data_dir}")
    print(f"   Processed output: {processed_dir}")
    print(f"   Chunk size: {chunk_size} characters")
    print(f"   Chunk overlap: {chunk_overlap} characters")
    print(f"   Reset vector DB: {reset_vectordb}")
    print("\n" + "="*80 + "\n")
    
    # -------------------------------------------------------------------------
    # STEP 1: DOCUMENT INGESTION
    # -------------------------------------------------------------------------
    print("📖 STEP 1/4: DOCUMENT INGESTION")
    print("-" * 80)
    print("Purpose: Extract text from PDF and DOCX files")
    print("Method: PyPDF2 for PDFs (no poppler needed), python-docx for DOCX")
    print()
    
    ingester = DocumentIngester(raw_data_dir)
    documents = ingester.process_all_documents()
    
    if not documents:
        print("\n❌ No documents processed! Check:")
        print(f"   1. Files exist in: {raw_data_dir}")
        print(f"   2. Files are valid PDF or DOCX format")
        print(f"   3. Files contain extractable text")
        return 0
    
    # Display statistics
    total_chars = sum(doc['char_count'] for doc in documents)
    avg_chars = total_chars // len(documents)
    
    print(f"\n✅ Successfully read {len(documents)} documents")
    print(f"   Total characters: {total_chars:,}")
    print(f"   Average per document: {avg_chars:,} characters")
    print(f"   File types: {sum(1 for d in documents if d['file_type']=='pdf')} PDFs, "
          f"{sum(1 for d in documents if d['file_type']=='docx')} DOCX")
    
    print("\n" + "="*80 + "\n")
    
    # -------------------------------------------------------------------------
    # STEP 2: TEXT CLEANING
    # -------------------------------------------------------------------------
    print("🧹 STEP 2/4: TEXT CLEANING")
    print("-" * 80)
    print("Purpose: Remove noise, fix formatting, normalize text")
    print("Operations:")
    print("   - Remove excessive whitespace")
    print("   - Fix hyphenated words across line breaks")
    print("   - Normalize spacing")
    print()
    
    cleaner = TextCleaner()
    
    for i, doc in enumerate(documents, 1):
        print(f"   Cleaning {i}/{len(documents)}: {doc['filename']}")
        doc['text'] = cleaner.clean(doc['text'])
    
    # Save cleaned documents
    cleaned_dir = Path(processed_dir) / "cleaned_text"
    cleaned_dir.mkdir(parents=True, exist_ok=True)
    
    for doc in documents:
        output_file = cleaned_dir / f"{Path(doc['filename']).stem}.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(doc['text'])
    
    print(f"\n✅ Text cleaned and saved to: {cleaned_dir}")
    print("\n" + "="*80 + "\n")
    
    # -------------------------------------------------------------------------
    # STEP 3: TEXT CHUNKING
    # -------------------------------------------------------------------------
    print("✂️  STEP 3/4: TEXT CHUNKING")
    print("-" * 80)
    print("Purpose: Split large documents into smaller, manageable chunks")
    print(f"Chunk size: {chunk_size} characters (max per chunk)")
    print(f"Overlap: {chunk_overlap} characters (preserves context)")
    print()
    print("Why chunking?")
    print("   - LLMs have token limits")
    print("   - Smaller chunks = more precise retrieval")
    print("   - Overlap ensures no context loss at boundaries")
    print()
    
    chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = chunker.chunk_documents(documents)
    
    if not chunks:
        print("\n❌ No chunks created!")
        return 0
    
    # Calculate statistics
    chunk_sizes = [len(chunk['text']) for chunk in chunks]
    avg_chunk_size = sum(chunk_sizes) // len(chunk_sizes)
    
    print(f"✅ Created {len(chunks)} chunks from {len(documents)} documents")
    print(f"   Average chunk size: {avg_chunk_size} characters")
    print(f"   Smallest chunk: {min(chunk_sizes)} characters")
    print(f"   Largest chunk: {max(chunk_sizes)} characters")
    
    # Save chunks
    chunks_dir = Path(processed_dir) / "chunks"
    chunker.save_chunks(chunks, str(chunks_dir))
    
    print("\n" + "="*80 + "\n")
    
    # -------------------------------------------------------------------------
    # STEP 4: EMBEDDING GENERATION & VECTOR STORAGE
    # -------------------------------------------------------------------------
    print("🧠 STEP 4/4: EMBEDDING GENERATION & VECTOR STORAGE")
    print("-" * 80)
    print("Purpose: Convert text chunks to vector embeddings for semantic search")
    print("Model: nomic-embed-text (768 dimensions)")
    print("Storage: ChromaDB vector database")
    print()
    print("What happens:")
    print("   1. Each chunk is converted to a 768-dimensional vector")
    print("   2. Vectors capture semantic meaning of the text")
    print("   3. Similar chunks will have similar vectors")
    print("   4. Enables fast similarity search during retrieval")
    print()
    
    embedding_manager = EmbeddingManager()
    collection = embedding_manager.store_chunks(chunks, reset_collection=reset_vectordb)
    
    print("\n" + "="*80)
    print("🎉 PIPELINE COMPLETE!")
    print("="*80)
    
    # Final summary
    print(f"\n📊 FINAL SUMMARY:")
    print(f"   ✅ Documents processed: {len(documents)}")
    print(f"   ✅ Chunks created: {len(chunks)}")
    print(f"   ✅ Embeddings generated: {len(chunks)}")
    print(f"   ✅ Vector database: data/vectordb/")
    print(f"   ✅ Total vectors in DB: {collection.count()}")
    
    print(f"\n💡 WHAT YOU CAN DO NOW:")
    print(f"   1. Query your documents:")
    print(f"      python -c \"from src.rag.rag_system import SiddhaRAG; rag = SiddhaRAG(); print(rag.query('your question'))\"")
    print(f"\n   2. Test retrieval:")
    print(f"      python src/rag/retriever.py")
    print(f"\n   3. Interactive Q&A:")
    print(f"      python src/rag/interactive_rag.py")
    
    print("\n" + "="*80 + "\n")
    
    return len(chunks)


# =============================================================================
# COMMAND LINE INTERFACE
# =============================================================================

def main():
    """
    Main entry point with command-line argument parsing.
    
    Usage:
        python src/pipeline/main_pipeline.py
        python src/pipeline/main_pipeline.py --reset-db
        python src/pipeline/main_pipeline.py --chunk-size 500
    """
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Siddha RAG Pipeline - Process documents and create vector database",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Run with default settings
  python src/pipeline/main_pipeline.py
  
  # Use smaller chunks
  python src/pipeline/main_pipeline.py --chunk-size 500 --chunk-overlap 100
  
  # Reset and rebuild database
  python src/pipeline/main_pipeline.py --reset-db
  
  # Specify custom data directory
  python src/pipeline/main_pipeline.py --raw-dir "path/to/documents"
        """
    )
    
    parser.add_argument(
        "--raw-dir",
        default="data/raw",
        help="Directory containing raw PDF/DOCX files (default: data/raw)"
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=1000,
        help="Maximum characters per chunk (default: 1000)"
    )
    parser.add_argument(
        "--chunk-overlap",
        type=int,
        default=200,
        help="Overlap between chunks in characters (default: 200)"
    )
    parser.add_argument(
        "--reset-db",
        action="store_true",
        help="Reset vector database before storing new data"
    )
    
    args = parser.parse_args()
    
    # Run pipeline with provided arguments
    try:
        chunks_created = run_complete_pipeline(
            raw_data_dir=args.raw_dir,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
            reset_vectordb=args.reset_db
        )
        
        if chunks_created > 0:
            print(f"✅ Success! Processed {chunks_created} chunks.")
            sys.exit(0)
        else:
            print("❌ Pipeline failed - no chunks created.")
            sys.exit(1)
            
    except Exception as e:
        print(f"\n❌ Pipeline error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
